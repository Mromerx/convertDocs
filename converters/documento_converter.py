import os
from pathlib import Path
from .libreoffice_engine import convert_with_libreoffice

SUPPORTED_FORMATS = ['docx', 'odt', 'pdf', 'txt']

def is_supported(ext: str):
    return ext.lower() in SUPPORTED_FORMATS

def convert(input_path: str, output_format: str, output_dir: str):
    """
    Converts a document to the specified format.
    """
    input_ext = Path(input_path).suffix.lower()[1:]
    output_format = output_format.lower()
    
    if input_ext == output_format:
        raise ValueError("Input and output format are the same.")
        
    if not is_supported(input_ext) or not is_supported(output_format):
        raise ValueError(f"Format not supported in document category. (Supported: {', '.join(SUPPORTED_FORMATS)})")
    
    output_file_name = f"{Path(input_path).stem}.{output_format}"
    output_path = os.path.join(output_dir, output_file_name)
    Path(output_dir).mkdir(parents=True, exist_ok=True)
    
    # Text Extraction: DOCX -> TXT
    if input_ext == 'docx' and output_format == 'txt':
        import docx
        doc = docx.Document(input_path)
        with open(output_path, 'w', encoding='utf-8') as f:
            for para in doc.paragraphs:
                f.write(para.text + "\n")
        return output_path
        
    # Text Generation: TXT -> DOCX
    elif input_ext == 'txt' and output_format == 'docx':
        import docx
        doc = docx.Document()
        with open(input_path, 'r', encoding='utf-8') as f:
            for line in f:
                doc.add_paragraph(line.strip("\n"))
        doc.save(output_path)
        return output_path
        
    # Special TXT <-> ODT cases (we route them through DOCX or use LibreOffice directly)
    # LibreOffice handles txt to odt and odt to txt natively very well.
    elif input_ext in ['docx', 'odt', 'txt', 'pdf'] and output_format in ['docx', 'odt', 'pdf', 'txt']:
        # If we reach here, we use LibreOffice engine as the fallback for all other valid combinations
        # e.g. DOCX->ODT, DOCX->PDF, ODT->PDF, ODT->DOCX, ODT->TXT, TXT->ODT, TXT->PDF
        return convert_with_libreoffice(input_path, output_format, output_dir)
        
    else:
        # For PDF -> TXT or TXT -> PDF, they shouldn't be handled here directly if it's the generic case,
        # but if a user asks to convert PDF to TXT in this module, we throw an error pointing to the generic converter
        if input_ext == 'pdf' and output_format == 'txt':
            raise ValueError("Use the generic PDF <-> TXT conversion")
            
        raise ValueError(f"Conversion {input_ext} -> {output_format} not implemented here.")
